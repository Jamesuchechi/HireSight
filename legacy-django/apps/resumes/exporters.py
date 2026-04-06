"""
Resume export functionality
Supports PDF and DOCX formats with template rendering
"""

import logging
from io import BytesIO
from typing import Optional, Dict, Any
from django.conf import settings

logger = logging.getLogger(__name__)

# PDF generation
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False
    logger.warning("WeasyPrint not installed. PDF export unavailable. Run: pip install weasyprint")

# DOCX generation
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX export unavailable. Run: pip install python-docx")


class PDFResumeExporter:
    """Export resume as PDF using WeasyPrint"""
    
    def __init__(self, resume, template_renderer):
        """
        Initialize PDF exporter
        
        Args:
            resume: Resume model instance
            template_renderer: ResumeTemplateRenderer instance
        """
        self.resume = resume
        self.renderer = template_renderer
    
    def export(self) -> BytesIO:
        """
        Generate PDF from resume
        
        Returns:
            BytesIO buffer containing PDF data
        """
        if not WEASYPRINT_AVAILABLE:
            raise ImportError("WeasyPrint is required for PDF export. Install with: pip install weasyprint")
        
        try:
            # Render HTML
            resume_data = self._prepare_resume_data()
            html_content = self.renderer.render_html(resume_data)
            
            # Add CSS
            css_content = self._get_css()
            
            # Generate PDF
            pdf_buffer = BytesIO()
            HTML(string=html_content).write_pdf(
                pdf_buffer,
                stylesheets=[CSS(string=css_content)]
            )
            pdf_buffer.seek(0)
            
            logger.info(f"Successfully generated PDF for resume: {self.resume.title}")
            return pdf_buffer
            
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            raise
    
    def _prepare_resume_data(self) -> Dict[str, Any]:
        """Prepare resume data for template"""
        parsed_data = self.resume.parsed_data or {}
        
        return {
            'title': self.resume.title,
            'user': {
                'name': self.resume.user.get_full_name(),
                'email': self.resume.user.email,
            },
            'parsed_data': parsed_data,
            'skills': parsed_data.get('skills', []),
            'education': parsed_data.get('education', []),
            'contact_info': parsed_data.get('contact_info', {}),
            'experience': parsed_data.get('experience', []),
            'certifications': parsed_data.get('certifications', []),
            'summary': parsed_data.get('summary', ''),
            'projects': parsed_data.get('projects', []),
        }
    
    def _get_css(self) -> str:
        """Get CSS for PDF rendering"""
        template_css = self.renderer.template.css_styles
        
        # Add print-specific CSS
        print_css = """
        @page {
            size: letter;
            margin: 0.5in;
        }
        body {
            font-family: 'Open Sans', Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            color: #333;
        }
        * {
            box-sizing: border-box;
        }
        """
        
        return template_css + print_css


class DOCXResumeExporter:
    """Export resume as DOCX using python-docx"""
    
    def __init__(self, resume, template_customization=None):
        """
        Initialize DOCX exporter
        
        Args:
            resume: Resume model instance
            template_customization: ResumeTemplateCustomization instance (optional)
        """
        self.resume = resume
        self.customization = template_customization
    
    def export(self) -> BytesIO:
        """
        Generate DOCX from resume
        
        Returns:
            BytesIO buffer containing DOCX data
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
        
        try:
            doc = Document()
            
            # Set margins
            self._set_margins(doc)
            
            # Add content
            self._add_header(doc)
            self._add_summary(doc)
            self._add_experience(doc)
            self._add_education(doc)
            self._add_skills(doc)
            self._add_certifications(doc)
            self._add_projects(doc)
            
            # Save to buffer
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            logger.info(f"Successfully generated DOCX for resume: {self.resume.title}")
            return docx_buffer
            
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            raise
    
    def _set_margins(self, doc):
        """Set document margins"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
    
    def _add_header(self, doc):
        """Add resume header with contact info"""
        parsed_data = self.resume.parsed_data or {}
        contact = parsed_data.get('contact_info', {})
        
        # Name
        name = contact.get('name', self.resume.user.get_full_name())
        heading = doc.add_heading(name, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_parts = []
        if contact.get('email'):
            contact_parts.append(contact['email'])
        if contact.get('phone'):
            contact_parts.append(contact['phone'])
        if contact.get('linkedin'):
            contact_parts.append(contact['linkedin'])
        if contact.get('github'):
            contact_parts.append(contact['github'])
        
        if contact_parts:
            p = doc.add_paragraph(' | '.join(contact_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Add spacing
        doc.add_paragraph()
    
    def _add_summary(self, doc):
        """Add professional summary"""
        parsed_data = self.resume.parsed_data or {}
        summary = parsed_data.get('summary', '')
        
        if summary:
            doc.add_heading('Professional Summary', level=2)
            doc.add_paragraph(summary)
            doc.add_paragraph()
    
    def _add_experience(self, doc):
        """Add work experience"""
        parsed_data = self.resume.parsed_data or {}
        experience = parsed_data.get('experience', [])
        
        if experience:
            doc.add_heading('Experience', level=2)
            for exp in experience:
                # Job title and company
                p = doc.add_paragraph()
                run = p.add_run(f"{exp.get('title', 'Position')} - {exp.get('company', 'Company')}")
                run.bold = True
                
                # Dates
                if exp.get('dates'):
                    doc.add_paragraph(exp['dates'], style='List Bullet')
                
                # Description
                if exp.get('description'):
                    doc.add_paragraph(exp['description'], style='List Bullet')
                
                # Add spacing between jobs
                doc.add_paragraph()
    
    def _add_education(self, doc):
        """Add education"""
        parsed_data = self.resume.parsed_data or {}
        education = parsed_data.get('education', [])
        
        if education:
            doc.add_heading('Education', level=2)
            for edu in education:
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                year = edu.get('year', '')
                
                text = f"{degree} - {institution}"
                if year:
                    text += f" ({year})"
                
                doc.add_paragraph(text, style='List Bullet')
            
            doc.add_paragraph()
    
    def _add_skills(self, doc):
        """Add skills"""
        parsed_data = self.resume.parsed_data or {}
        skills = parsed_data.get('skills', [])
        
        if skills:
            doc.add_heading('Skills', level=2)
            doc.add_paragraph(', '.join(skills))
            doc.add_paragraph()
    
    def _add_certifications(self, doc):
        """Add certifications"""
        parsed_data = self.resume.parsed_data or {}
        certifications = parsed_data.get('certifications', [])
        
        if certifications:
            doc.add_heading('Certifications', level=2)
            for cert in certifications:
                cert_text = cert.get('text', cert) if isinstance(cert, dict) else cert
                doc.add_paragraph(cert_text, style='List Bullet')
            
            doc.add_paragraph()
    
    def _add_projects(self, doc):
        """Add projects"""
        parsed_data = self.resume.parsed_data or {}
        projects = parsed_data.get('projects', [])
        
        if projects:
            doc.add_heading('Projects', level=2)
            for project in projects:
                # Project name
                p = doc.add_paragraph()
                run = p.add_run(project.get('name', 'Project'))
                run.bold = True
                
                # Description
                if project.get('description'):
                    doc.add_paragraph(project['description'], style='List Bullet')
                
                # Technologies
                if project.get('technologies'):
                    tech_text = f"Technologies: {', '.join(project['technologies'])}"
                    doc.add_paragraph(tech_text, style='List Bullet')
                
                doc.add_paragraph()


class PlainTextExporter:
    """Export resume as plain text"""
    
    def __init__(self, resume):
        self.resume = resume
    
    def export(self) -> str:
        """
        Generate plain text from resume
        
        Returns:
            Plain text string
        """
        return self.resume.parsed_text or "No content available"
