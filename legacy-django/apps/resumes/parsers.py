import re
import spacy
from typing import Dict, List, Optional, Tuple
from django.conf import settings
from django.core.files.storage import default_storage
import fitz  # PyMuPDF for PDF parsing
from docx import Document  # python-docx for DOCX parsing
import logging
import os
import tempfile

logger = logging.getLogger(__name__)


class ResumeParserError(Exception):
    """Custom exception for resume parsing errors."""
    pass


class CertificationExtractor:
    """Extract certifications from resume text."""

    # Ordered by priority (highest accuracy first)
    PATTERNS = [
        {
            'name': 'Industry Standard',
            'regex': r'(?:^|\b)((?:AWS|Azure|GCP|Google|Cisco|CompTIA|Microsoft|Oracle|Salesforce|VMware|Linux|ITIL|PMI|ISC|GIAC)[A-Za-z0-9\s,&\-]*?(?:Certified|Certification|Certificate|Cert|Credential|License)(?:\s+(?:Associate|Professional|Expert|Developer|Architect|Administrator|Security|Cloud|Solutions|Data|Practitioner|Engineer))?)\b',
            'priority': 1,
            'confidence': 0.92,
        },
        {
            'name': 'Trademarked Programs',
            'regex': r'(?:^|\b)((?:Salesforce|ServiceNow|Atlassian|HubSpot|Marketo|Tableau|Looker)\s+[A-Za-z\s\-]*?(?:Certification|Certified|Badge|Credential)(?:\s+(?:Administrator|Developer|Consultant|Expert))?)\b',
            'priority': 2,
            'confidence': 0.88,
        },
        {
            'name': 'Educational Platform',
            'regex': r'(?:^|\b)((?:Certified|Certificate)\s+(?:in|for|as)?\s+[A-Z][A-Za-z\s\-&]+(?:Program|Course|Track|Specialization|Bootcamp)?)\b',
            'priority': 3,
            'confidence': 0.85,
        },
        {
            'name': 'Professional License',
            'regex': r'(?:^|\b)((?:License|License[d]|Licensure|Accreditation)\s+(?:as|in|for)?\s*(?:[A-Z]{2,}|[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)(?:\s*(?:#|No\.|Number)?\s*[0-9]+)*)\b',
            'priority': 4,
            'confidence': 0.88,
        },
        {
            'name': 'Military Clearance',
            'regex': r'(?:^|\b)((?:Security\s+Clearance|Top\s+Secret|Secret|Confidential|TS/SCI|TS/SCI\s+with\s+(?:CI\s+)?Poly))\b',
            'priority': 5,
            'confidence': 0.95,
        },
    ]

    # Blacklist: Common false positives to exclude
    BLACKLIST_PATTERNS = [
        r'certified\s+(?:mail|check|copy)',  # Postal terms
        r'certified\s+(?:organic|fair\s+trade)',  # Product labels
        r'software\s+certification',  # Too generic
        r'driver\'?s?\s+license',  # Already captured by license pattern
    ]

    def extract_certifications(self, text: str) -> List[Tuple[str, float, str]]:
        """
        Extract certifications from resume text.

        Args:
            text: Resume text

        Returns:
            List of (certification, confidence_score, pattern_type)
        """
        certifications = []
        found_certs = set()  # Track duplicates

        # Apply each pattern in order
        for pattern_obj in self.PATTERNS:
            matches = re.finditer(
                pattern_obj['regex'],
                text,
                re.IGNORECASE | re.MULTILINE
            )

            for match in matches:
                cert_text = match.group(1).strip()

                # Skip if blacklisted
                if self._is_blacklisted(cert_text):
                    continue

                # Skip if already found (deduplicate)
                cert_normalized = cert_text.lower()
                if cert_normalized in found_certs:
                    continue

                # Calculate confidence (pattern confidence × text quality score)
                confidence = pattern_obj['confidence'] * self._quality_score(cert_text)

                certifications.append({
                    'text': cert_text,
                    'confidence': confidence,
                    'pattern': pattern_obj['name'],
                    'priority': pattern_obj['priority'],
                })

                found_certs.add(cert_normalized)

        # Sort by priority and confidence
        certifications.sort(
            key=lambda x: (x['priority'], -x['confidence'])
        )

        return certifications

    def _is_blacklisted(self, text: str) -> bool:
        """Check if text matches blacklist patterns."""
        for pattern in self.BLACKLIST_PATTERNS:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False

    def _quality_score(self, text: str) -> float:
        """
        Score text quality (0-1).
        Better formatted/longer certs score higher.
        """
        length = len(text)

        # Ideal length: 15-80 characters
        if 15 <= length <= 80:
            quality = 1.0
        elif 10 <= length < 15 or 80 < length <= 120:
            quality = 0.85
        else:
            quality = 0.6

        # Boost if contains uppercase words (proper formatting)
        uppercase_count = sum(1 for c in text if c.isupper())
        if uppercase_count >= 3:
            quality = min(1.0, quality + 0.1)

        return quality


class ResumeParser:
    """AI-powered resume parser using spaCy."""

    def __init__(self):
        # Load spaCy model
        try:
            self.nlp = spacy.load("en_core_web_sm")
            logger.info("spaCy model loaded successfully")
        except OSError:
            logger.warning("spaCy model not available, using fallback parser")
            self.nlp = None

        # Initialize certification extractor
        self.cert_extractor = CertificationExtractor()

    def parse_content(self, file_content: bytes, filename: str) -> Dict:
        """
        Parse resume from file content (bytes) and extract structured data.

        Args:
            file_content: Binary content of the uploaded file
            filename: Original filename (to determine file type)

        Returns:
            Dict containing parsed data with 'success' key
        """
        try:
            # Extract text from content
            text = self._extract_text_from_content(file_content, filename)

            if not text or len(text.strip()) < 50:
                return {
                    'success': False,
                    'error': 'Could not extract sufficient text from file. File may be empty or corrupted.'
                }

            # Parse text with NLP
            parsed_data = self._parse_text(text)

            return {
                'success': True,
                'text': text,
                **parsed_data
            }

        except Exception as e:
            logger.error(f"Error parsing resume {filename}: {str(e)}", exc_info=True)
            return {
                'success': False,
                'error': f'Parsing error: {str(e)}'
            }

    def parse_file(self, file_path: str, filename: str) -> Dict:
        """
        Parse resume file and extract structured data.

        Args:
            file_path: Path to the uploaded file (can be relative or absolute)
            filename: Original filename

        Returns:
            Dict containing parsed data with 'success' key
        """
        try:
            actual_path = self._get_file_path(file_path)
            # Extract text from file
            text = self._extract_text_from_file(actual_path, filename)

            if not text or len(text.strip()) < 50:
                return {
                    'success': False,
                    'error': 'Could not extract sufficient text from file. File may be empty or corrupted.'
                }

            # Parse text with NLP
            parsed_data = self._parse_text(text)

            return {
                'success': True,
                'text': text,
                **parsed_data
            }

        except Exception as e:
            logger.error(f"Error parsing resume {filename}: {str(e)}", exc_info=True)
            return {
                'success': False,
                'error': f'Parsing error: {str(e)}'
            }

    def _extract_text_from_content(self, file_content: bytes, filename: str) -> str:
        """Extract text from file content (bytes)."""
        file_extension = filename.lower().split('.')[-1]

        try:
            if file_extension == 'pdf':
                return self._extract_pdf_text_from_content(file_content)
            elif file_extension in ['docx', 'doc']:
                return self._extract_docx_text_from_content(file_content)
            elif file_extension == 'txt':
                return file_content.decode('utf-8', errors='ignore').strip()
            else:
                raise ResumeParserError(f"Unsupported file type: {file_extension}")

        except ResumeParserError:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise ResumeParserError(f"Could not extract text: {str(e)}")

        except ResumeParserError:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise ResumeParserError(f"Could not extract text: {str(e)}")

    def _extract_text_from_file(self, file_path: str, filename: str) -> str:
        """Extract text from file based on its type."""
        file_extension = filename.lower().split('.')[-1]
        
        try:
            if file_extension == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_extension in ['docx', 'doc']:
                return self._extract_docx_text(file_path)
            elif file_extension == 'txt':
                return self._extract_txt_text(file_path)
            else:
                raise ResumeParserError(f"Unsupported file type: {file_extension}")
        
        except ResumeParserError:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise ResumeParserError(f"Could not extract text: {str(e)}")

    def _get_file_path(self, file_path: str) -> str:
        """
        Get the actual file path, handling both relative and absolute paths.
        If using Django storage, download the file to a temporary location.
        """
        # Check if it's already an absolute path and exists
        if os.path.isabs(file_path) and os.path.exists(file_path):
            logger.info(f"Using absolute path: {file_path}")
            return file_path
        
        # Check if relative path exists from current directory
        if os.path.exists(file_path):
            abs_path = os.path.abspath(file_path)
            logger.info(f"Converted relative to absolute: {abs_path}")
            return abs_path
        
        # If using Django's default storage (S3, local media, etc.)
        try:
            if default_storage.exists(file_path):
                logger.info(f"File exists in Django storage: {file_path}")
                
                # Check if we're using FileSystemStorage
                if hasattr(default_storage, 'path'):
                    # Local file system
                    actual_path = default_storage.path(file_path)
                    logger.info(f"Using filesystem storage path: {actual_path}")
                    return actual_path
                else:
                    # Remote storage (S3, etc.) - download to temp file
                    logger.info(f"Downloading from remote storage: {file_path}")
                    with default_storage.open(file_path, 'rb') as remote_file:
                        # Create a temporary file
                        suffix = os.path.splitext(file_path)[1]
                        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                        temp_file.write(remote_file.read())
                        temp_file.close()
                        logger.info(f"Downloaded to temp file: {temp_file.name}")
                        return temp_file.name
        except Exception as e:
            logger.error(f"Error accessing file via Django storage: {e}")
        
        # Last resort: check in MEDIA_ROOT
        if hasattr(settings, 'MEDIA_ROOT'):
            media_path = os.path.join(settings.MEDIA_ROOT, file_path)
            if os.path.exists(media_path):
                logger.info(f"Found in MEDIA_ROOT: {media_path}")
                return media_path
        
        # If we get here, file cannot be found
        raise ResumeParserError(
            f"File not found: {file_path}. "
            f"Checked: absolute path, relative path, Django storage, and MEDIA_ROOT."
        )

    def _extract_text_from_file(self, file_path: str, filename: str) -> str:
        """Extract text from file based on its type."""
        file_extension = filename.lower().split('.')[-1]
        
        try:
            if file_extension == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_extension in ['docx', 'doc']:
                return self._extract_docx_text(file_path)
            elif file_extension == 'txt':
                return self._extract_txt_text(file_path)
            else:
                raise ResumeParserError(f"Unsupported file type: {file_extension}")
        
        except ResumeParserError:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise ResumeParserError(f"Could not extract text: {str(e)}")

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        temp_file_created = False
        
        try:
            # Check if this is a temp file we created
            if file_path.startswith(tempfile.gettempdir()):
                temp_file_created = True
             
            with fitz.open(file_path) as doc:
                for page_num, page in enumerate(doc, 1):
                    page_text = page.get_text()
                    if page_text:
                        text += page_text + "\n"
                     
                    # Safety check - limit to reasonable number of pages
                    if page_num > 20:
                        logger.warning(f"PDF has more than 20 pages, stopping at page {page_num}")
                        break
             
            return text.strip()

        except Exception as e:
            raise ResumeParserError(f"PDF extraction failed: {str(e)}")
         
        finally:
            # Clean up temp file if we created one
            if temp_file_created and os.path.exists(file_path):
                try:
                    os.unlink(file_path)
                    logger.info(f"Cleaned up temp file: {file_path}")
                except Exception as e:
                    logger.warning(f"Could not delete temp file {file_path}: {e}")

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        temp_file_created = file_path.startswith(tempfile.gettempdir())
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts)

        except Exception as e:
            raise ResumeParserError(f"DOCX extraction failed: {str(e)}")
        
        finally:
            # Clean up temp file if we created one
            if temp_file_created and os.path.exists(file_path):
                try:
                    os.unlink(file_path)
                    logger.info(f"Cleaned up temp file: {file_path}")
                except Exception as e:
                    logger.warning(f"Could not delete temp file {file_path}: {e}")

    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file."""
        temp_file_created = file_path.startswith(tempfile.gettempdir())
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return text.strip()

        except Exception as e:
            raise ResumeParserError(f"TXT extraction failed: {str(e)}")
        
        finally:
            # Clean up temp file if we created one
            if temp_file_created and os.path.exists(file_path):
                try:
                    os.unlink(file_path)
                    logger.info(f"Cleaned up temp file: {file_path}")
                except Exception as e:
                    logger.warning(f"Could not delete temp file {file_path}: {e}")

    def _extract_pdf_text_from_content(self, file_content: bytes) -> str:
        """Extract text from PDF file content (bytes)."""
        text = ""
        
        try:
            # Create temporary file from bytes
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False, mode='wb') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
             
            try:
                with fitz.open(temp_file_path) as doc:
                    for page_num, page in enumerate(doc, 1):
                        page_text = page.get_text()
                        if page_text:
                            text += page_text + "\n"
                         
                        # Safety check - limit to reasonable number of pages
                        if page_num > 20:
                            logger.warning(f"PDF has more than 20 pages, stopping at page {page_num}")
                            break
                 
                return text.strip()
             
            finally:
                # Clean up temp file
                if os.path.exists(temp_file_path):
                    try:
                        os.unlink(temp_file_path)
                    except Exception as e:
                        logger.warning(f"Could not delete temp PDF file: {e}")
         
        except Exception as e:
            raise ResumeParserError(f"PDF extraction failed: {str(e)}")

    def _extract_docx_text_from_content(self, file_content: bytes) -> str:
        """Extract text from DOCX file content (bytes)."""
        try:
            # Create temporary file from bytes
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False, mode='wb') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
             
            try:
                doc = Document(temp_file_path)
                text_parts = []
                 
                # Extract from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                 
                # Extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)
                 
                return "\n".join(text_parts)
     
            finally:
                # Clean up temp file
                if os.path.exists(temp_file_path):
                    try:
                        os.unlink(temp_file_path)
                    except Exception as e:
                        logger.warning(f"Could not delete temp DOCX file: {e}")
    
        except Exception as e:
            raise ResumeParserError(f"DOCX extraction failed: {str(e)}")

    def _parse_text(self, text: str) -> Dict:
        """Parse text using NLP to extract structured data."""
        if not self.nlp:
            return self._parse_text_fallback(text)

        try:
            # Process text with spaCy (limit to first 1 million characters)
            doc = self.nlp(text[:1000000])

            # Extract various components
            skills = self._extract_skills(doc, text)
            experience_years = self._extract_experience_years(text)
            education = self._extract_education(doc, text)
            contact_info = self._extract_contact_info(text)
            certifications = self.cert_extractor.extract_certifications(text)

            return {
                'skills': skills,
                'experience_years': experience_years,
                'education': education,
                'contact_info': contact_info,
                'certifications': certifications,
            }

        except Exception as e:
            logger.error(f"NLP parsing failed: {str(e)}, falling back to basic parser")
            return self._parse_text_fallback(text)

    def _extract_skills(self, doc, text: str) -> List[str]:
        """Extract skills from parsed document."""
        skills = set()

        # Expanded tech skills dictionary
        tech_skills = {
            # Programming languages
            'python', 'javascript', 'java', 'c++', 'c#', 'ruby', 'php', 'go', 'rust',
            'swift', 'kotlin', 'typescript', 'scala', 'r', 'matlab', 'perl', 'dart',
            
            # Frontend
            'react', 'angular', 'vue', 'svelte', 'next.js', 'nuxt', 'redux', 'jquery',
            'html', 'html5', 'css', 'css3', 'sass', 'less', 'tailwind', 'bootstrap',
            'material-ui', 'chakra', 'webpack', 'vite', 'babel',
            
            # Backend
            'django', 'flask', 'fastapi', 'express', 'nodejs', 'node.js', 'spring',
            'laravel', 'rails', 'asp.net', '.net', 'graphql', 'rest', 'grpc',
            
            # Databases
            'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
            'cassandra', 'dynamodb', 'sqlite', 'oracle', 'mariadb', 'neo4j',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'ansible',
            'jenkins', 'gitlab', 'github', 'circleci', 'travis', 'nginx', 'apache',
            
            # Data & AI
            'machine learning', 'deep learning', 'ai', 'nlp', 'computer vision',
            'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy',
            'spark', 'hadoop', 'tableau', 'power bi', 'jupyter',
            
            # Tools
            'git', 'linux', 'unix', 'bash', 'powershell', 'vim', 'vscode', 'jira',
            'confluence', 'agile', 'scrum', 'kanban',
        }

        # Extract from skills section (if exists)
        skills_section = self._extract_section(text, ['skills', 'technical skills', 'competencies'])
        if skills_section:
            for skill in tech_skills:
                if skill.lower() in skills_section.lower():
                    skills.add(skill.title())

        # Extract from entire text
        text_lower = text.lower()
        for skill in tech_skills:
            if skill in text_lower:
                skills.add(skill.title())

        # Extract nouns and proper nouns as potential skills
        for token in doc:
            if token.pos_ in ['NOUN', 'PROPN'] and len(token.text) > 2:
                skill = token.text.lower()
                if skill in tech_skills or self._is_technical_skill(skill):
                    skills.add(token.text)

        # Remove duplicates and sort
        return sorted(list(skills))[:50]  # Limit to 50 skills

    def _extract_section(self, text: str, headers: List[str]) -> str:
        """Extract a section from resume based on header keywords."""
        text_lower = text.lower()
        
        for header in headers:
            pattern = rf'\n\s*{re.escape(header)}\s*\n(.*?)(?=\n\s*[A-Z][^a-z]*\n|$)'
            match = re.search(pattern, text_lower, re.DOTALL | re.IGNORECASE)
            if match:
                return match.group(1)
        
        return ""

    def _is_technical_skill(self, word: str) -> bool:
        """Check if a word looks like a technical skill."""
        tech_indicators = [
            'api', 'sdk', 'ide', 'cli', 'ui', 'ux', 'ci/cd',
            'dev', 'ops', 'web', 'mobile', 'cloud', 'data',
            'code', 'script', 'framework', 'library', 'tool'
        ]
        return any(indicator in word for indicator in tech_indicators)

    def _extract_experience_years(self, text: str) -> Optional[float]:
        """Extract years of experience from text."""
        patterns = [
            r'(\d+(?:\.\d+)?)\s*(?:\+\s*)?years?\s*(?:of\s*)?experience',
            r'experience\s*(?:of\s*)?(\d+(?:\.\d+)?)\s*(?:\+\s*)?years?',
            r'(\d+(?:\.\d+)?)\+?\s*yrs?\s*(?:of\s*)?(?:exp|experience)',
        ]

        years_found = []
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            years_found.extend([float(match) for match in matches])

        if years_found:
            # Take the highest number found (likely total experience)
            return max(years_found)

        return None

    def _extract_education(self, doc, text: str) -> List[Dict]:
        """Extract education information."""
        education = []

        # Common degree patterns
        degree_patterns = [
            r'(ph\.?d\.?|doctorate|doctor of philosophy)',
            r'(master\'?s?|m\.?s\.?|m\.?a\.?|mba|master of)',
            r'(bachelor\'?s?|b\.?s\.?|b\.?a\.?|b\.?tech|bachelor of)',
            r'(associate\'?s?|a\.?s\.?|a\.?a\.?|associate of)',
        ]

        # Education keywords
        education_keywords = [
            'university', 'college', 'institute', 'school',
            'bachelor', 'master', 'phd', 'degree', 'diploma',
            'certification', 'certificate'
        ]

        # Extract education section
        edu_section = self._extract_section(
            text,
            ['education', 'academic background', 'qualifications']
        )

        if edu_section:
            # Split by common separators
            entries = re.split(r'\n{2,}|\n\s*[-•]\s*', edu_section)
            
            for entry in entries:
                if any(keyword in entry.lower() for keyword in education_keywords):
                    education.append({
                        'text': entry.strip(),
                        'degree': self._extract_degree(entry),
                        'institution': self._extract_institution(entry),
                        'year': self._extract_year(entry)
                    })

        return education[:5]  # Limit to 5 education entries

    def _extract_degree(self, text: str) -> str:
        """Extract degree from education text."""
        degrees = {
            'PhD': r'ph\.?d\.?|doctorate|doctor of philosophy',
            'Master': r'master\'?s?|m\.?s\.?|m\.?a\.?|mba|master of',
            'Bachelor': r'bachelor\'?s?|b\.?s\.?|b\.?a\.?|b\.?tech|bachelor of',
            'Associate': r'associate\'?s?|a\.?s\.?|a\.?a\.?|associate of',
        }
        
        for degree, pattern in degrees.items():
            if re.search(pattern, text, re.IGNORECASE):
                return degree
        
        return ""

    def _extract_institution(self, text: str) -> str:
        """Extract institution name from education text."""
        # Look for university/college/institute
        institution_pattern = r'(?:at\s+|from\s+)?([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:University|College|Institute))'
        match = re.search(institution_pattern, text)
        
        if match:
            return match.group(1)
        
        return ""

    def _extract_year(self, text: str) -> Optional[int]:
        """Extract graduation year."""
        year_pattern = r'\b(19|20)\d{2}\b'
        matches = re.findall(year_pattern, text)
        
        if matches:
            years = [int(y) for y in matches]
            # Return most recent year
            return max(years)
        
        return None

    def _extract_contact_info(self, text: str) -> Dict:
        """Extract contact information."""
        contact_info = {}

        # Email regex (improved)
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info['email'] = emails[0]

        # Phone regex (multiple formats)
        phone_patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
        ]
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                contact_info['phone'] = phones[0]
                break

        # LinkedIn URL
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
        linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin:
            contact_info['linkedin'] = linkedin[0]

        # GitHub URL
        github_pattern = r'(?:https?://)?(?:www\.)?github\.com/[\w-]+'
        github = re.findall(github_pattern, text, re.IGNORECASE)
        if github:
            contact_info['github'] = github[0]

        return contact_info

    def _parse_text_fallback(self, text: str) -> Dict:
        """Fallback parsing without spaCy."""
        return {
            'skills': self._extract_skills_fallback(text),
            'experience_years': self._extract_experience_years(text),
            'education': [],
            'contact_info': self._extract_contact_info(text),
            'certifications': self.cert_extractor.extract_certifications(text),
        }

    def _extract_skills_fallback(self, text: str) -> List[str]:
        """Fallback skill extraction without NLP."""
        skills = set()
        
        common_skills = [
            'python', 'javascript', 'java', 'react', 'sql',
            'aws', 'docker', 'git', 'html', 'css'
        ]
        
        text_lower = text.lower()
        for skill in common_skills:
            if skill in text_lower:
                skills.add(skill.title())
        
        return sorted(list(skills))


# Global parser instance
resume_parser = ResumeParser()
